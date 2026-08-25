# -*- coding: utf-8 -*-
"""Genera el informe en un archivo Word (.docx) a partir del dict que arma
informe.generar(), incluyendo una gráfica de evolución histórica."""

import io
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

AZUL = RGBColor(0x2A, 0x78, 0xD6)
GRIS = RGBColor(0x52, 0x51, 0x4E)


def _grafica_precio(precios) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(6.3, 3), dpi=150)
    ax.plot(precios.index, precios.values, color="#2a78d6", linewidth=1.6)
    if len(precios) >= 200:
        sma = precios.rolling(200).mean()
        ax.plot(sma.index, sma.values, color="#eb6834", linewidth=1.2,
               linestyle="--", label="Media móvil 200 días")
        ax.legend(fontsize=8, frameon=False)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=8, colors="#52514e")
    ax.set_ylabel("Precio", fontsize=9, color="#52514e")
    ax.grid(axis="y", color="#e1e0d9", linewidth=0.6)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


def _titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


_PATRON_NEGRITA = re.compile(r"\*\*(.+?)\*\*")


def _parrafo(doc, texto, negrita=False, tam=10.5, color=None):
    """Agrega un párrafo, convirtiendo el **negrita** estilo Markdown del
    catálogo educativo en negrita real de Word (no asteriscos literales)."""
    p = doc.add_paragraph()
    pos = 0
    for m in _PATRON_NEGRITA.finditer(texto):
        if m.start() > pos:
            r = p.add_run(texto[pos:m.start()])
            r.bold = negrita
            r.font.size = Pt(tam)
            if color:
                r.font.color.rgb = color
        r = p.add_run(m.group(1))
        r.bold = True
        r.font.size = Pt(tam)
        if color:
            r.font.color.rgb = color
        pos = m.end()
    if pos < len(texto):
        r = p.add_run(texto[pos:])
        r.bold = negrita
        r.font.size = Pt(tam)
        if color:
            r.font.color.rgb = color
    return p


def _tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = h
        celda.paragraphs[0].runs[0].bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val)
    return t


def _fmt_usd(v):
    if v is None:
        return "—"
    if abs(v) >= 1e9:
        return f"US$ {v/1e9:,.1f} mil millones"
    if abs(v) >= 1e6:
        return f"US$ {v/1e6:,.1f} millones"
    return f"US$ {v:,.0f}"


def _fmt_pct(v, dec=1):
    return "—" if v is None else f"{v*100:.{dec}f}%"


def generar_docx(r: dict) -> io.BytesIO:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = GRIS

    # ── portada ──
    t = doc.add_heading(f"{r['nombre']} ({r['ticker']})", level=0)
    for run in t.runs:
        run.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    sub.add_run(
        f"Informe de activo · {r['categoria'] or 'Sin categoría'} · "
        f"Precio actual: {r['precio_actual']:,.2f} {r['moneda']} · "
        f"Generado el {r['fecha_informe'].strftime('%d/%m/%Y')}"
    ).italic = True

    _parrafo(doc,
        "⚠️ Este informe es una herramienta educativa para aprender sobre "
        "inversiones. No constituye asesoramiento financiero ni una "
        "recomendación de compra o venta.", negrita=True)

    # ── qué es ──
    _titulo(doc, "1. ¿Qué es y cómo funciona?")
    edu = r["educativo"]
    if edu["intro_categoria"]:
        _parrafo(doc, edu["intro_categoria"])
    _parrafo(doc, edu["que_es"])
    if edu["como_funciona"]:
        _parrafo(doc, "Cómo gana dinero: ", negrita=True)
        _parrafo(doc, edu["como_funciona"])

    # ── wikipedia ──
    if r["wiki"]:
        _titulo(doc, "2. Contexto independiente (Wikipedia)", nivel=2)
        _parrafo(doc, r["wiki"]["extracto"])
        _parrafo(doc, f"Fuente: Wikipedia — {r['wiki']['titulo']} "
                     f"({r['wiki']['url']})", tam=8.5)

    # ── desempeño ──
    _titulo(doc, "3. Desempeño de precio")
    rend = r["rendimiento"]
    _tabla(doc, ["Período", "Rendimiento"], [
        ["En lo que va del año (YTD)", _fmt_pct(rend["ytd"])],
        ["Último año", _fmt_pct(rend["un_anio"])],
        ["Anualizado a 5 años (CAGR)", _fmt_pct(rend["cagr_5y"])],
    ])

    doc.add_paragraph()
    _titulo(doc, "Evolución histórica (5 años)", nivel=2)
    img = _grafica_precio(r["precios"])
    doc.add_picture(img, width=Inches(6.0))

    # ── cuánto gana la empresa (SEC) ──
    if r["sec"]:
        _titulo(doc, "4. Cuánto gana la empresa (fuente oficial: SEC EDGAR)")
        _parrafo(doc,
            "Estos son los datos que la propia empresa presentó ante el "
            "regulador bursátil de EE.UU. (SEC) en su último balance anual "
            "(10-K) — una fuente independiente de Yahoo Finance.")
        sec = r["sec"]
        filas = []
        if sec["ingresos_valor"]:
            filas.append([f"Ingresos (año fiscal {sec['ingresos_anio_fiscal']})",
                          _fmt_usd(sec["ingresos_valor"])])
        if sec["ganancia_valor"]:
            filas.append([f"Ganancia neta (año fiscal {sec['ganancia_anio_fiscal']})",
                          _fmt_usd(sec["ganancia_valor"])])
        if filas:
            _tabla(doc, ["Concepto", "Monto"], filas)
        _parrafo(doc, f"Fuente: SEC EDGAR ({sec['url']})", tam=8.5)

    # ── indicadores ──
    _titulo(doc, "5. Indicadores clave")
    fu = r["fundamentales"]
    ri = r["riesgo"]
    _tabla(doc, ["Indicador", "Valor"], [
        ["Volatilidad anual", _fmt_pct(ri["volatilidad"])],
        ["Peor caída (drawdown máximo)", _fmt_pct(ri["max_drawdown"])],
        ["Ratio de Sharpe", "—" if ri["sharpe"] is None else f"{ri['sharpe']:.2f}"],
        ["Dividendo anual (yield)", _fmt_pct(fu["div_yield"]) if fu["div_yield"] else "No paga"],
        ["P/E (precio/ganancias)", "—" if not fu["pe"] else f"{fu['pe']:.1f}"],
        ["Beta (vs. mercado)", "—" if not fu["beta"] else f"{fu['beta']:.2f}"],
        ["Capitalización de mercado", _fmt_usd(fu["market_cap"])],
    ])

    # ── impuestos ──
    _titulo(doc, "6. Impuestos para un inversor uruguayo")
    imp = r["impuesto"]
    _parrafo(doc, "Sobre dividendos/intereses:", negrita=True)
    _parrafo(doc, f"Retención en origen (EE.UU.): "
                 f"{_fmt_pct(imp['origen_tasa'])}", negrita=True)
    _parrafo(doc, imp["origen_motivo"])
    _parrafo(doc, f"Carga fiscal total combinada (con IRPF uruguayo): "
                 f"{_fmt_pct(imp['carga_total'])}", negrita=True)
    _parrafo(doc, imp["uy_nota"])
    _parrafo(doc, "Sobre la ganancia de capital al vender:", negrita=True)
    _parrafo(doc, f"IRPF sobre la ganancia de capital: "
                 f"{_fmt_pct(imp['gc_tasa'])}", negrita=True)
    _parrafo(doc, imp["gc_nota"])

    # ── riesgos y perfil ──
    if edu["riesgos"]:
        _titulo(doc, "7. Riesgos principales")
        _parrafo(doc, edu["riesgos"])
    if edu["perfil"]:
        _titulo(doc, "8. ¿Para qué perfil de inversor?")
        _parrafo(doc, edu["perfil"])

    # ── fuentes consultadas ──
    _titulo(doc, "Fuentes consultadas")
    fuentes_txt = [f"Yahoo Finance — precios, dividendos y datos "
                   f"fundamentales (consultado el {r['fecha_consulta']})"]
    if r["wiki"]:
        fuentes_txt.append(f"Wikipedia — {r['wiki']['titulo']}")
    if r["sec"]:
        fuentes_txt.append("SEC EDGAR — balance anual oficial (10-K)")
    for f in fuentes_txt:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph()
    _parrafo(doc,
        "Este informe fue generado automáticamente con fines educativos. "
        "La información puede contener errores o estar desactualizada: "
        "verificá siempre los datos antes de tomar una decisión de "
        "inversión, y consultá a un profesional matriculado si necesitás "
        "asesoramiento financiero o tributario personalizado.",
        tam=8.5)

    salida = io.BytesIO()
    doc.save(salida)
    salida.seek(0)
    return salida
